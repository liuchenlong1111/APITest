import json
import re
from typing import Optional, Dict, Any
from pathlib import Path
import aiofiles
import docx
import openpyxl
import markdown
from bs4 import BeautifulSoup


class DocumentParser:
    """文档解析器"""
    
    @staticmethod
    async def parse_document(file_path: str, file_type: str = None) -> Dict[str, Any]:
        """
        解析文档内容
        
        Args:
            file_path: 文件路径
            file_type: 文件类型，如果为None则从文件扩展名推断
        
        Returns:
            包含文档内容和元数据的字典
        """
        if file_type is None:
            file_type = Path(file_path).suffix.lower()
        
        try:
            if file_type in ['.txt', '.md', '.markdown']:
                return await DocumentParser._parse_text_file(file_path, file_type)
            elif file_type == '.docx':
                return await DocumentParser._parse_docx_file(file_path)
            elif file_type in ['.xlsx', '.xls']:
                return await DocumentParser._parse_excel_file(file_path)
            elif file_type in ['.html', '.htm']:
                return await DocumentParser._parse_html_file(file_path)
            elif file_type == '.json':
                return await DocumentParser._parse_json_file(file_path)
            else:
                # 尝试作为文本文件处理
                return await DocumentParser._parse_text_file(file_path, '.txt')
        
        except Exception as e:
            return {
                "content": "",
                "type": file_type,
                "error": f"解析文件失败: {str(e)}",
                "metadata": {}
            }
    
    @staticmethod
    async def _parse_text_file(file_path: str, file_type: str) -> Dict[str, Any]:
        """解析文本文件"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
            content = await file.read()
        
        # 如果是Markdown文件，转换为HTML并提取纯文本
        if file_type in ['.md', '.markdown']:
            html = markdown.markdown(content, extensions=['extra', 'codehilite'])
            soup = BeautifulSoup(html, 'html.parser')
            plain_text = soup.get_text()
            
            return {
                "content": content,
                "plain_text": plain_text,
                "html": html,
                "type": "markdown",
                "metadata": {
                    "line_count": len(content.split('\n')),
                    "char_count": len(content)
                }
            }
        
        return {
            "content": content,
            "type": "text",
            "metadata": {
                "line_count": len(content.split('\n')),
                "char_count": len(content)
            }
        }
    
    @staticmethod
    async def _parse_docx_file(file_path: str) -> Dict[str, Any]:
        """解析Word文档"""
        doc = docx.Document(file_path)
        
        # 提取段落文本
        paragraphs = []
        tables = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append({
                    "text": paragraph.text,
                    "style": paragraph.style.name if paragraph.style else None
                })
        
        # 提取表格
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            tables.append(table_data)
        
        # 组合所有文本内容
        all_text = '\n'.join([p["text"] for p in paragraphs])
        
        # 添加表格文本
        for table in tables:
            for row in table:
                all_text += '\n' + '\t'.join(row)
        
        return {
            "content": all_text,
            "type": "docx",
            "paragraphs": paragraphs,
            "tables": tables,
            "metadata": {
                "paragraph_count": len(paragraphs),
                "table_count": len(tables),
                "char_count": len(all_text)
            }
        }
    
    @staticmethod
    async def _parse_excel_file(file_path: str) -> Dict[str, Any]:
        """解析Excel文件"""
        workbook = openpyxl.load_workbook(file_path, data_only=True)
        sheets_data = {}
        all_text = ""
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sheet_data = []
            
            for row in sheet.iter_rows(values_only=True):
                row_data = []
                for cell in row:
                    cell_value = str(cell) if cell is not None else ""
                    row_data.append(cell_value)
                    all_text += cell_value + "\t"
                sheet_data.append(row_data)
                all_text += "\n"
            
            sheets_data[sheet_name] = sheet_data
        
        return {
            "content": all_text,
            "type": "excel",
            "sheets": sheets_data,
            "metadata": {
                "sheet_count": len(sheets_data),
                "char_count": len(all_text)
            }
        }
    
    @staticmethod
    async def _parse_html_file(file_path: str) -> Dict[str, Any]:
        """解析HTML文件"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
            html_content = await file.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # 提取纯文本
        plain_text = soup.get_text()
        
        # 提取所有链接
        links = []
        for a in soup.find_all('a', href=True):
            links.append({
                "text": a.get_text().strip(),
                "href": a['href']
            })
        
        # 提取表格
        tables = []
        for table in soup.find_all('table'):
            table_data = []
            for row in table.find_all('tr'):
                row_data = []
                for cell in row.find_all(['td', 'th']):
                    row_data.append(cell.get_text().strip())
                if row_data:
                    table_data.append(row_data)
            if table_data:
                tables.append(table_data)
        
        return {
            "content": plain_text,
            "html": html_content,
            "type": "html",
            "links": links,
            "tables": tables,
            "metadata": {
                "link_count": len(links),
                "table_count": len(tables),
                "char_count": len(plain_text)
            }
        }
    
    @staticmethod
    async def _parse_json_file(file_path: str) -> Dict[str, Any]:
        """解析JSON文件"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
            json_content = await file.read()
        
        try:
            data = json.loads(json_content)
            
            # 如果是OpenAPI规范
            if DocumentParser._is_openapi_spec(data):
                return DocumentParser._parse_openapi_spec(data)
            
            # 如果是Postman集合
            elif DocumentParser._is_postman_collection(data):
                return DocumentParser._parse_postman_collection(data)
            
            # 普通JSON文件
            else:
                return {
                    "content": json.dumps(data, indent=2, ensure_ascii=False),
                    "type": "json",
                    "data": data,
                    "metadata": {
                        "keys": list(data.keys()) if isinstance(data, dict) else [],
                        "char_count": len(json_content)
                    }
                }
        
        except json.JSONDecodeError as e:
            return {
                "content": json_content,
                "type": "json",
                "error": f"JSON格式错误: {str(e)}",
                "metadata": {}
            }
    
    @staticmethod
    def _is_openapi_spec(data: Dict) -> bool:
        """判断是否为OpenAPI规范"""
        return (
            isinstance(data, dict) and 
            ('openapi' in data or 'swagger' in data) and 
            'paths' in data
        )
    
    @staticmethod
    def _is_postman_collection(data: Dict) -> bool:
        """判断是否为Postman集合"""
        return (
            isinstance(data, dict) and 
            'info' in data and 
            'item' in data and 
            isinstance(data.get('info'), dict) and 
            'name' in data['info']
        )
    
    @staticmethod
    def _parse_openapi_spec(data: Dict) -> Dict[str, Any]:
        """解析OpenAPI规范"""
        spec_version = data.get('openapi', data.get('swagger', 'unknown'))
        info = data.get('info', {})
        paths = data.get('paths', {})
        
        # 提取API信息
        apis = []
        for path, methods in paths.items():
            for method, details in methods.items():
                if method.lower() in ['get', 'post', 'put', 'delete', 'patch', 'head', 'options']:
                    api_info = {
                        "path": path,
                        "method": method.upper(),
                        "summary": details.get('summary', ''),
                        "description": details.get('description', ''),
                        "tags": details.get('tags', []),
                        "parameters": details.get('parameters', []),
                        "responses": details.get('responses', {}),
                        "security": details.get('security', [])
                    }
                    apis.append(api_info)
        
        content = f"""
OpenAPI规范文档解析结果:
版本: {spec_version}
标题: {info.get('title', '未知')}
版本: {info.get('version', '未知')}
描述: {info.get('description', '无描述')}

包含 {len(apis)} 个API接口:
""" + "\n".join([f"- {api['method']} {api['path']}: {api['summary']}" for api in apis])
        
        return {
            "content": content,
            "type": "openapi",
            "spec_version": spec_version,
            "info": info,
            "apis": apis,
            "metadata": {
                "api_count": len(apis),
                "version": spec_version
            }
        }
    
    @staticmethod
    def _parse_postman_collection(data: Dict) -> Dict[str, Any]:
        """解析Postman集合"""
        info = data.get('info', {})
        items = data.get('item', [])
        
        apis = []
        
        def extract_requests(items, prefix=""):
            for item in items:
                if 'request' in item:
                    request = item['request']
                    method = request.get('method', 'GET')
                    url = request.get('url', {})
                    
                    if isinstance(url, str):
                        path = url
                    elif isinstance(url, dict):
                        path = url.get('raw', '')
                    else:
                        path = str(url)
                    
                    api_info = {
                        "name": item.get('name', ''),
                        "method": method,
                        "path": path,
                        "description": item.get('description', ''),
                        "headers": request.get('header', []),
                        "body": request.get('body', {}),
                        "auth": request.get('auth', {})
                    }
                    apis.append(api_info)
                
                # 递归处理子项目
                if 'item' in item:
                    extract_requests(item['item'], prefix + item.get('name', '') + "/")
        
        extract_requests(items)
        
        content = f"""
Postman集合解析结果:
名称: {info.get('name', '未知')}
描述: {info.get('description', '无描述')}

包含 {len(apis)} 个API请求:
""" + "\n".join([f"- {api['method']} {api['name']}: {api['path']}" for api in apis])
        
        return {
            "content": content,
            "type": "postman",
            "info": info,
            "apis": apis,
            "metadata": {
                "api_count": len(apis),
                "collection_name": info.get('name', '未知')
            }
        }


# 全局实例
document_parser = DocumentParser() 